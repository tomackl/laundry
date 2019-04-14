from docx import Document
from docx.shared import Inches
import pandas as pd
import janitor
from pathlib import Path
from typing import List, Iterable, Dict, Tuple, Any, NewType


data_frame = NewType('data_frame', pd.DataFrame)


def clean_xlsx_table(file_path: str, sheet: str, head: int = 0,
                     rm_column: List[str] = None, clean_hdr: bool = False,
                     drop_empty: bool = False) -> data_frame:
    """
    Open and perform basic data cleaning on a single excel work sheet.
    :param file_path: path to the excel file
    :param sheet: excel spreadsheet name
    :param head: index of the header row in the spreadsheet. Defaults to 0.
    :param rm_column: remove the column headers contained in the passed list.
    :param clean_hdr: If True clean the column headers
    :param drop_empty: If True remove empty rows.
    :return:
    """
    df = pd.read_excel(pd.ExcelFile(file_path), sheet, head)
    if rm_column is not None:
        df = df.remove_columns(rm_column)
    if clean_hdr is not False:
        df = df.clean_names()
    if drop_empty is True:
        df = df.dropna()
    return df


def extract_data(record: Dict, header: List[str], format_title: bool = True
                 ) -> List[Tuple]:
    """
    Take a dictionary and split in to a list of tuples containing the 'keys'
    data defined in 'header' as the first tuple and the associated values as the second.
    The function will return a list containing two equal length tuples.
    :param record: the dictionary containing the data.
    :param header: the keys that defined the key-values to be extracted.
    :param format_title: make the header string title case.
    :return: a list of tuples
    """
    hdr, data = [], []
    for each in header:
        h = remove_underscore(each)
        if format_title is True:
            hdr.append(str(h).title())
        else:
            hdr.append(str(h))
        t = record.pop(each)
        data.append(t)
    return [tuple(hdr), tuple(data)]


def remove_underscore(text: str) -> str:
    return text.replace('_', ' ')


def insert_paragraph(document: object, text: str, title: str = None,
                     section_style: str = None, title_style: str = None,
                     ):
    """
    This function defines the characteristics of a paragraph to be added to the document.
    :param document: the template_doc the paragraph will be added to.
    :param text: paragraph text
    :param title: title text
    :param section_style: paragraph style
    :param title_style: title style. Use this _or_ title_level.
    :return:
    """
    if (title is not None) and (title_style is not None):
        document.add_paragraph(title, style=title_style)
    document.add_paragraph(str(text), style=section_style)


def insert_table(document: object, cols: int, rows: int,
                 data: List[Iterable[str]], section_style: str = None,
                 autofit: bool = True):
    """
    The function takes data and uses it to create a table for the template_doc.
    The first row of data is assumed to be the table header.
    :param document: the docx file the table will be added to.
    :param rows: the number of required table rows.
    :param cols: the number of required table columns.
    :param data: The list data to be inserted into the table. The idx[0] is assumed to be the header.
    :param section_style: The style to be used for the table.
    :param autofit: autofit the table to the page width.
    :return:
    """
    # todo: add error checking for 'cols' and 'row'
    table: object = document.add_table(rows=rows, cols=cols, style=section_style)
    if autofit is True:
        table.autofit = True
    data = enumerate(data, 0)
    for i, cell_contents in data:
        insert_row(table.rows[i].cells, cell_contents)


def insert_row(row_cells, data: List[str]) -> Any:
    """
    Populate a table row. The cells are passed as a row and the contents added.
    :param row_cells:
    :param data:
    :param style: style is the text style to be applied to the individual rows.
    :return:
    """
    for i, text in enumerate(data):
        row_cells[i].text = str(text)
    return row_cells


def insert_photo(document: object, photo: str, width: int = 4):
    """
    Insert a photo located at path into document and set the photo to width.
    :param document: the docx file the photo will be added to.
    :param photo: file path to the
    :param width: width of the image in Inches
    :return:
    """
    # todo: add ability to have more than one photo
    if photo.lower() != 'no photo':
        photo_path = Path(photo)
        document.add_picture(str(photo_path), width=Inches(width))

def format_docx(rowdict: dict, structdict: dict, outputfile: object):
    """
    The function is passed a dict (data_dict) containing the data to be formatted
    (structure) based on the template (outputfile).
    :param rowdict: dictionary containing the text. It represents a single row from the spreadsheet.
    :param structdict: defines the output file's format structure.
    :param outputfile: The file which data will be inserted into.
    :return:
    """
    # todo: some though needs to be given to how to format the output titles.
    #       consideration to the use of .strip() has been given however this might cause
    #       problems for users.

    # todo: add error checking here.
    for element in structdict:
        if str(element['sectiontype']).lower() in ('heading', 'para', 'paragraph'):
            insert_paragraph(outputfile, str(rowdict[element['sectioncontains']]),
                             title=str(element['sectioncontains']).title(),
                             section_style=element['sectionstyle'],
                             title_style=element['titlestyle']
                             )

        elif str(element['sectiontype']).lower() == 'table':
            sect_contains = []
            # the below assumes that the headers are divided by new lines or by commas.
            # todo: convert this section into a function. this should be the same the photo seciton below (common function)
            if '\n' in element['sectioncontains']:
                sect_contains = list(str(element['sectioncontains']).splitlines())
            elif ',' in element['sectioncontains']:
                sect_contains = list(str(element['sectioncontains']).split(','))
            data = extract_data(rowdict, sect_contains)
            insert_table(outputfile, len(sect_contains), len(data),
                         data, section_style=element['sectionstyle']
                         )

        # todo: this section should refernce the coloumn title in the spreadsheet structure work sheet.
        # todo: find a way of inserting the path and the file extension into the structure work sheet
        elif str(element['sectiontype']).lower() == 'photo':
            # insert_photo(outputfile, rowdict['location'], 4)
            if str(rowdict['photo']).lower() in ['nan']:
                break
            if '\n' in rowdict['photo']:
                rowdict['photo'] = list(str(rowdict['photo']).splitlines())
            else:
                rowdict['photo'] = list(str(rowdict['photo']).split(','))
            for each in rowdict['photo']:
                loc = str(rowdict['location']) \
                      + each \
                      + str(rowdict['file_extension'])
                insert_photo(outputfile, loc, 4)

        else:
            print('Valid section header was not found.')

        if element['sectionbreak'] is True:
            insert_paragraph(outputfile, '')

        if element['pagebreak'] is True:
            outputfile.add_page_break()

# ==> variables <==


"""
the following has been implemented:

* sectioncontains: list[str] -> containing the column titles
* sectiontype: str -> paragraph/table/heading/photo
* sectionstyle: str -> Word paragraph style
* titlestyle: str -> Word title style. This does not apply to tables.
* sectionbreak: bool -> is a break required after each section?
* pagebreak: bool -> is a page break required after the section?
    - => Headings should be a single column per paragraph
    - => Paragraphs should be a single column per paragraph
    - => Tables can be any number of columns (there will be practical limits)
    - => section_styles _can_ only be a single value
    - => title_styles _can_ only be a single value
    - => pagebreak is a True/False value
"""

# ==> cli arguments <==
# todo: The expected cli form is templar input_file output_file template_file

# todo: cli options include -s --structure -> the worksheet containing the document structure
# todo: the path needs to provided to click
dir = '../../resources/input_files/'
file = 'test_spreadsheet.xlsm'
path = dir+file
data_worksheet = 'Master List'
structure_worksheet = '_structure_'


# TODO: provide an option to provide a path to the templating template_doc.
template_doc = Document('../../resources/templates/template.docx')

# ==> data cleanup arguments <==
# todo: provide a way of describing the columns that need to be removed from the spreadsheet
remove_columns = []
#     'Recommended Actions',
#     'Comment',
#     'Link',
# ]Link

# todo: define a way of replacing columns names with new names.
#       the list below isn't actually implemented within the script.
new_cols = [(
    ' p&s ID',
    'Hazard_ID'
)]

# todo: the following information needs to be provided ... somewhere/somehow
#       - path
#       - worksheet
#       - first row of data (head)
#       - which columns are to be removed
#       - are headers to be cleaned?
#       - are empty columns to be dropped.
#       This information should be separated from the general formatting requirements.

# IMPORT THE DATA FROM THE SPREADSHEET
data_file = clean_xlsx_table(path, sheet=data_worksheet, head=5,
                             rm_column=remove_columns, clean_hdr=True,
                             drop_empty=False
                             )
data_dict = data_file.to_dict('records')

# IMPORT THE OUTPUT DOCUMENT STRUCTURE FROM THE SPREADSHEET
structure_file = clean_xlsx_table(path, sheet=structure_worksheet, head=0,
                                  clean_hdr=True, drop_empty=False
                                  )
structure_dict = structure_file.to_dict('records')

for row in data_dict:
    format_docx(row, structure_dict, template_doc)
template_doc.save('../../resources/output_files/converted_file.docx')
