"""Main class for laundry. This is intended to replace the original laundry script."""

from laundry.constants import data_frame, invalid, photo_formats
from typing import Dict, List, Iterable, Tuple, NamedTuple, Any
from docx import Document
from docx.shared import Inches
from pathlib import Path, PurePath
import janitor
import pandas as pd
from colorama import init as colorama_init
from colorama import Fore, Back, Style
from sys import exit as sys_exit

colorama_init(autoreset=True)

# Define headers for the batch and structure worksheets. These are fixed.
EXPECTED_BATCH_HEADERS = ['data_worksheet', 'structure_worksheet', 'header_row', 'drop_empty_columns', 'template_file',
                          'filter_rows', 'output_file']
EXPECTED_STRUCTURE_HEADERS = ['section_type', 'section_contains', 'section_style', 'title_style', 'section_break',
                              'page_break', 'path']
EXPECTED_SECTION_TYPES = ['heading', 'table', 'para', 'photo']
OUTPUT_TITLE = {'fore_colour': 'GREEN', 'style_colour': 'BRIGHT'}
OUTPUT_TEXT = {'fore_colour': 'GREEN', 'style_colour': 'DIM'}
EXCEPTION_TEXT = {'fore_colour': 'RED', 'style_colour': 'BRIGHT'}
DATAFRAME_TITLE = {'fore_colour': 'BLUE', 'style_colour': 'BRIGHT'}
DATAFRAME_TEXT = {'fore_colour': 'BLACK', 'back_colour': 'BLUE'}
FAULTFIND_TEXT = {'fore_colour': 'BLACK', 'back_colour': 'GREEN'}
OUTPUT_SUCCESS = {'fore_colour': 'CYAN', 'back_colour': 'BLACK', 'style_colour': 'BRIGHT'}


def exit_app():
    sys_exit()


def print_verbose(text: (str, Exception), verbose: bool, fore_colour: str = 'RESET', back_colour: str = 'RESET',
                  style_colour: str = 'NORMAL', end: str = '\n', flush: bool = True):
    """
    A function to provide clearer informaiton to the user at the command line.
    :param text: Text to be printed
    :param verbose: Is the text to be output.
    :param fore_colour: The text's colour.
    :param back_colour: The background colour.
    :param style_colour: The text style.
    :param end:
    :param flush:
    :return:
    """
    if verbose is False:
        return
    fore = {'BLACK': Fore.BLACK, 'RED': Fore.RED, 'GREEN': Fore.GREEN, 'YELLOW': Fore.YELLOW, 'BLUE': Fore.BLUE,
            'MAGENTA': Fore.MAGENTA, 'CYAN': Fore.CYAN, 'WHITE': Fore.WHITE, 'RESET': Fore.RESET}

    back = {'BLACK': Back.BLACK, 'RED': Back.RED, 'GREEN': Back.GREEN, 'YELLOW': Back.YELLOW, 'BLUE': Back.BLUE,
            'MAGENTA': Back.MAGENTA, 'CYAN': Back.CYAN, 'WHITE': Back.WHITE, 'RESET': Back.RESET}

    style = {'DIM': Style.DIM, 'NORMAL': Style.NORMAL, 'BRIGHT': Style.BRIGHT, 'RESET_ALL': Style.RESET_ALL}

    if fore_colour is not None:
        fore = fore[fore_colour.upper()]
    if back_colour is not None:
        back = back[back_colour.upper()]
    if style_colour is not None:
        style = style[style_colour.upper()]

    print(f'{fore}{back}{style}{text}{Style.RESET_ALL}', end=end, flush=flush)


def split_str(data_str: str) -> List[str]:
    """
    Split a string into a list of strings. The string will be split at ',' or '\n' with '\n' taking precedence.
    :param data_str: A string to be split.
    :return: List[str]
    """
    if '\n' in data_str:
        return list(str(data_str).splitlines())
    elif ',' in data_str:
        return list(str(data_str).split(','))
    else:
        i = list()
        i.append(data_str)
        return i


def sort_table_data(record: Dict, header: List[str], format_title: bool = True) -> List[Tuple]:
    """
    Take a dictionary and split in to a list of tuples containing the 'keys' cell_data defined in 'header' as the
    first tuple and the associated values as the second. The function will return a list containing two equal length
    tuples. The function assumes that the cell_data contained in record and header is in the correct order.
    :param record: The dictionary containing the cell_data.
    :param header: The keys that defined the key-values to be extracted.
    :param format_title: If True make the header string title case.
    :return: A list of tuples containing the header and cell_data information.
    """
    hdr_list: List[str] = []
    data_list: List[str] = []
    for each in header:
        hdr_data = remove_underscore(each)
        if format_title is True:
            hdr_data = hdr_data.title()
        hdr_list.append(hdr_data)
        data_list.append(record.pop(each))
    return [tuple(hdr_list), tuple(data_list)]


def remove_underscore(data_str: str) -> str:
    """
    A function to remove underscores from a string returning the string with spaces instead.
    :param data_str:
    :return:
    """
    return data_str.replace('_', ' ')


def resolve_file_path(path: (Path, str)) -> (Path, Exception):
    """
    Resolve the passed filepath returning as a Path() or an exception.
    :return:
    """
    q = str(path).replace('\\', '/')
    try:
        p = Path(q).resolve(strict=True)
        return p
    except Exception as e:
        print_verbose(f'{e}', True, **EXCEPTION_TEXT)


def values_exist(expected: set, actual: set) -> bool:
    """
    Check the expected worksheets exist.
    :return:
    """
    return expected <= actual


def remove_from_iterable(values: Iterable, *args) -> List:
    """Return items in list that are not equal to a value in drop."""
    data = []
    for each in values:
        if each not in args:
            data.append(each)
    return data


def strip_whitespace(wht_spc: List) -> List:
    i = []
    for each in wht_spc:
        i.append(each.strip())
    return i


class SingleLoad:
    """
    This class is intended to replace the original Laundry's procedural approach from the single load function.
    The click.progressbar() will not be transferred from the original function.
    Assumptions
    1. The file paths passed are correct and have already been checked.
    2. Data passed to the class is in the correct format.
    3. The output file will be created by the object.
    """

    def __init__(self, structure_data: pd.DataFrame, data_data: pd.DataFrame, file_template: Path,
                 file_output_path: Path):
        """
        # The method signature is based on the laundry.single_load() function. This calls self.format_docx()
        :param structure_data: A dictionary that defines the structure of the documentation.
        :param data_data: A dictionary that contains the cell_data to be formatted.
        :param file_template: The Word .docx file that contains the formatting styles to be used.
        :param file_output_path: The path to the output file location.
        """
        self._structure: pd.DataFrame = structure_data
        self._data: pd.DataFrame = data_data
        self._file_template: Document() = Document(file_template)
        self._file_output: Path = Path(file_output_path)
        self._row_data: List[Dict] = list()
        self.start_wash()
        self.issue_document()

    def start_wash(self):
        """
        Start formatting the output document.
        """
        for row in self._data.itertuples():
            self.format_docx(row)

        print_verbose(f'\nDocument {self._file_output} completed', True, **OUTPUT_SUCCESS)

    def format_docx(self, row: NamedTuple):
        """
        This is factory method that calls the appropriate the information contained within document structure.
        :param row: dictionary containing the data_str to be formatted. This is a single row from the spreadsheet.
        :return:
        """

        row = row._asdict()
        for structure_element in self._structure.itertuples():
            sect_contains_element: str = str(structure_element.section_contains).lower()
            sect_style_element: str = str(structure_element.section_style)
            sect_type_element: str = str(structure_element.section_type).lower()
            title_style_element: str = str(structure_element.title_style)
            sect_break_element: bool = structure_element.section_break
            page_break_element: bool = structure_element.page_break

            if sect_type_element in ['heading', 'para', 'paragraph']:
                # removed .lower() from the string passed to the insert_paragraph call
                self.insert_paragraph(str(row[sect_contains_element]), title=sect_contains_element.title(),
                                      section_style=sect_style_element, title_style=title_style_element)

            elif sect_type_element == 'table':
                table_col_hdr = split_str(sect_contains_element)
                sorted_row = sort_table_data(row, table_col_hdr)
                self.insert_table(len(table_col_hdr), len(sorted_row), sorted_row, section_style=sect_style_element)

            elif sect_type_element == 'photo':
                if isinstance(row[sect_contains_element], Iterable):
                    for each in row[sect_contains_element]:
                        self.insert_photo(each, 4)
            else:
                print('Valid section header was not found.')

            if sect_break_element is True:
                self.insert_paragraph('')

            if page_break_element is True:
                self._file_template.add_page_break()

    def insert_paragraph(self, text: str, title: str = None, section_style: str = None,
                         title_style: str = None):
        """
        This method defines the characteristics of a paragraph to be added to the document.
        :param text: paragraph data_str
        :param title: title data_str
        :param section_style: paragraph style
        :param title_style: title style. Use this _or_ title_level.
        :return:
        """
        split_text = text.splitlines()
        if len(split_text) == 0:
            self._file_template.add_paragraph(split_text)
        else:
            if (title is not None) and (str(title_style) != 'nan'):
                self._file_template.add_paragraph(remove_underscore(title), style=title_style)
            for each in split_text:
                self._file_template.add_paragraph(each, style=section_style)

    def insert_table(self, cols: int, rows: int, data: List[Iterable[str]], section_style: str = None,
                     autofit_table: bool = True):
        """
        The function takes data and uses it to create a table for the template_doc.
        The first row of data is assumed to be the table header.
        todo: add 'add_table_hdr added to allow the table header to be dropped from the table if not required.
        :param rows: the number of required table rows.
        :param cols: the number of required table columns.
        :param data: The list data to be inserted into the table. The idx[0] is assumed to be the header.
        :param section_style: The style to be used for the table.
        :param autofit_table: autofit the table to the page width.
        :return:
        """
        table = self._file_template.add_table(rows=rows, cols=cols, style=section_style)
        table.autofit = autofit_table
        cell_data = enumerate(data, 0)
        for i, cell_contents in cell_data:
            for j, text in enumerate(cell_contents):
                table.rows[i].cells[j].text = str(text)

    def insert_photo(self, photo: Path, width: int = 4):
        """
        Insert a photo located at path into document and set the photo width.
        :param photo: file path to the
        :param width: width of the image in Inches
        :return:
        """
        self._file_template.add_picture(str(photo), width=Inches(width))

    def issue_document(self):
        """
        Output the file.
        :return:
        """
        self._file_template.save(self._file_output)


class Laundry:
    def __init__(self, input_fp: Path, data_worksheet: str = None, structure_worksheet: str = None,
                 batch_worksheet: str = None, header_row: int = 0, drop_empty_columns: bool = None,
                 template_file: str = None, filter_rows: str = None, output_file: (Path, str) = None,
                 verbose: bool = True, template_generate: bool = False):
        """
        Instantiating the class will run error checking on the passed information, checking for the following steps:
        1. A basic check that worksheet names have been passed.
        2. The input file's path is correct.
        3. Check that the data, structure and batch worksheet names passed at instantiation exist within the file. This
            is the first worksheet name check, additional check will be made as required.
        4. If the batching information is passed to the object at instantiation, then merge this into a dictionary.
            Error checking will be completed later.
        5.
        :param input_fp: The file path to the spreadsheet containing the data
        :param data_worksheet: The name of the worksheet containing the data to be formatted.
        :param structure_worksheet: The name of the worksheet containing the output document's structure.
        :param batch_worksheet: The name of the worksheet containing the batch data.
        :param header_row: 
        :param drop_empty_columns: An explicit tag to drop empty rows from the worksheet if they contain two or more
        empty cells. If this is left as None it will be automatically set to True for the data worksheet.
        :param template_file: 
        :param filter_rows: 
        :param output_file:
        :param verbose:
        :param template_generate:
        """
        if template_generate:
            # Generate the template spreadsheet and exit the app.
            self.generate_tempate_document()

        self.output_verbose: bool = verbose
        # Step 1: Basic data checking.
        t_sheets_expected = remove_from_iterable([data_worksheet, structure_worksheet, batch_worksheet], None)
        print_verbose('Check: Worksheets are present:', verbose=self.output_verbose, **OUTPUT_TITLE)
        if len(t_sheets_expected) == 0:
            raise ValueError(f'Either the "data" and "structure" worksheets, or the "batch" worksheet must be '
                             f'provided.')
        sheet = enumerate(t_sheets_expected, 1)
        for item, sht in sheet:
            print_verbose(f'  Sheet {item}:\t{sht}', verbose=self.output_verbose, **OUTPUT_TEXT)

        # Step 2: Confirm the input file exists.
        self._input_fp: (Path, str) = ''
        try:
            print_verbose(f'Check: Resolving spreadsheet filepath: ', verbose=self.output_verbose, **OUTPUT_TITLE)
            self._input_fp = resolve_file_path(input_fp)
            print_verbose(f'  {self._input_fp}', verbose=self.output_verbose, **OUTPUT_TEXT)
        except Exception as e:
            print_verbose(f'\t{e}: File {input_fp} does not exist.', True, **EXCEPTION_TEXT)

        # Load the Excel file into memory.
        self._washing_basket = pd.ExcelFile(self._input_fp)

        # Gather the worksheet names
        self._sheets_actual: list = self._washing_basket.sheet_names

        # Set up the lists to contain the dictionaries containing: 1. data, 2. output structure, and 3. batch docs
        self._data: List[dict] = []
        self._structure: List[dict] = []
        self._batch: List[dict] = []

        #  Step 3: Check that the data, structure and batch worksheet names passed exist within the file.
        try:
            print_verbose(f'Check Worksheets exist in spreadsheet: ', verbose=self.output_verbose, **OUTPUT_TITLE)
            self.compare_lists(t_sheets_expected, self._sheets_actual)
            t_sht_actual = enumerate(self._sheets_actual, 1)
            for item, sht in t_sht_actual:
                print_verbose(f'  Sheet {item}:\t {sht}', verbose=self.output_verbose, **OUTPUT_TEXT)
        except Exception as e:
            print_verbose(f'{e}', True, **EXCEPTION_TEXT)

        # Step 4. If the batching information is passed to the object at instantiation, then merge this into a
        #   dictionary. Error checking will be completed later.
        # Step 4.1. If all the expected command line parameters are set to the defaults assume that the a batch
        #   approach has been used. We do _not_ test for batch since this will be tested for later.
        input_arg = [data_worksheet, structure_worksheet, header_row, drop_empty_columns, template_file, filter_rows,
                     output_file]

        self.batch_df: pd.DataFrame = pd.DataFrame(columns=['data_worksheet', 'structure_worksheet', 'header_row',
                                                            'drop_empty_columns', 'template_file', 'output_file'])
        # Step 4.2. Since the default values for the input args are all 'None' or 0, if we remove these values from the
        #   list, if the list's length is greater than 0 then there is a chance that a single wash is required. We don't
        #   test for the input file path since this has already occurred.
        if len(remove_from_iterable(input_arg, None, 0)) > 0:

            # If the drop_empty_rows is None set it to True. This will save the user problems.
            if drop_empty_columns is None:
                drop_empty_columns = True

            # Step 4.3. Turn the command line arguments into a dict and store temporarily.
            t_batch_dict = {'data_worksheet': [data_worksheet], 'structure_worksheet': [structure_worksheet],
                            'header_row': [header_row], 'drop_empty_columns': [drop_empty_columns],
                            'template_file': [template_file], 'filter_rows': [filter_rows],
                            'output_file': [output_file]}
            self.batch_df = pd.DataFrame.from_dict(data=t_batch_dict)
        # Step 5. If batch information passed as a worksheet clean and sort the batch data.
        else:
            self.batch_df = self.excel_to_dataframe(self._washing_basket, batch_worksheet, header_row=0,
                                                    clean_header=True)

        # Step 6. Check the batch data.
        try:
            print_verbose(f'Batch worksheet data', verbose=self.output_verbose, **DATAFRAME_TITLE)
            print_verbose(f'{self.batch_df}', verbose=self.output_verbose, **DATAFRAME_TEXT)
            print_verbose(f'Check: Batch worksheet data', verbose=self.output_verbose, **OUTPUT_TITLE)
            self.check_batch_worksheet_data()
            print_verbose(f'Batch data checked', verbose=self.output_verbose, **OUTPUT_TITLE)
        except Exception as e:
            print_verbose(f'{e}', True, **EXCEPTION_TEXT)
            exit_app()

        # Step 6. Convert the batch DataFrame to a dict and store.
        self._batch_dict = self.batch_df.to_dict('records')

        # Step 7 - Every row of the the batch DataFrame contains information regarding an output file. For each row in
        # the DataFrame produce the associated output file.
        for t_batch_row in self.batch_df.itertuples():
            t_structure_worksheet = t_batch_row.structure_worksheet
            t_data_worksheet = t_batch_row.data_worksheet
            self.t_structure_df = self.excel_to_dataframe(self._washing_basket, t_structure_worksheet, header_row=0,
                                                          clean_header=True, drop_empty_rows=False)

            self.t_structure_photo_path: Dict[str, Path] = {}
            self.t_data_df = self.excel_to_dataframe(self._washing_basket, t_data_worksheet,
                                                     header_row=t_batch_row.header_row,
                                                     clean_header=True, drop_empty_rows=True)

            # Filter the data DataFrame using the filters passed.
            if str(t_batch_row.filter_rows).lower() not in invalid and t_batch_row.filter_rows is not None:
                for row_filter in t_batch_row.filter_rows:
                    self.t_data_df = self.t_data_df.loc[self.t_data_df[row_filter[0]].isin(row_filter[1])]

            # Step 8 - Check the structure data.
            self.check_dataframe(f'Structure worksheet data', self.t_structure_df, f'Check: Structure worksheet data',
                                 self.check_structure_worksheet_data, f'Structure dataframe checked',
                                 f'{t_batch_row.structure_worksheet}', f'Structure dataframe failure: ')

            # Step 9 - Check the data worksheet data.
            self.check_dataframe(f'Data dataframe', self.t_data_df, f'Check: Data worksheet data',
                                 self.check_data_worksheet_data, f'Data dataframe checked',
                                 f'{t_batch_row.data_worksheet}', f'Data dataframe failure: ')

            self.wash_load(t_batch_row.template_file, t_batch_row.output_file)

            del self.t_structure_photo_path

    def generate_tempate_document(self):
        """
        Generate a blank teamplate.
        :return:
        """
        df_batch = pd.DataFrame(columns=EXPECTED_BATCH_HEADERS)
        series_section_type = {key: '' for key in EXPECTED_STRUCTURE_HEADERS}
        series_section_type['section_type'] = EXPECTED_SECTION_TYPES
        df_structure = pd.DataFrame.from_dict(series_section_type)
        with pd.ExcelWriter('Laundry_template.xlsx') as writer:
            df_batch.to_excel(writer, sheet_name='_batch', index=False)
            df_structure.to_excel(writer, sheet_name='_structure', index=False)
        print_verbose('Template file saved.', verbose=True, **OUTPUT_TEXT)
        exit_app()

    def wash_load(self, template_file: Path, output_file: Path):
        """

        :param template_file:
        :param output_file:
        :return:
        """
        SingleLoad(self.t_structure_df, self.t_data_df, template_file, output_file)

    def check_batch_worksheet_data(self):
        """
        Check the batch worksheet data is in the correct format. Data is checked as a DataFrame. The following checks
        are made.
        Check 1: Confirm expected batch headers exist in the batch worksheet.
        Check 2: Confirm Structure and data worksheets referenced in batch worksheet exist.
        Check 3: Confirm the template files exist and resolve the files.
        Check 4: Confirm an output filename has been provided.
        Check 5: Check if filter_rows
        Check 6: Check if drop_empty_rows is None, set it to False.
        Check 7: Check if header_row is None, set it to 0.
        :return:
        """
        # Extract the data and structure worksheet names from the batch worksheet for error checking.
        t_batch_headers = list(self.batch_df)
        t_batch_data_worksheets_expected = list(self.batch_df.loc[:, 'data_worksheet'])
        t_batch_structure_worksheets_expected = list(self.batch_df.loc[:, 'structure_worksheet'])

        # Check 1.
        self.data_check(f'\tBatch work sheet headers.', f'Ok', [(EXPECTED_BATCH_HEADERS, t_batch_headers)],
                        compare='subset')

        # Check 2.
        self.data_check(f'\tData & structure worksheets referenced correctly.', f'Ok',
                        [(t_batch_structure_worksheets_expected, self._sheets_actual),
                         (t_batch_data_worksheets_expected, self._sheets_actual)], compare='subset')

        # Check 3.
        for row in self.batch_df.itertuples():
            print_verbose(f'  Row {row.Index}:', verbose=self.output_verbose, **OUTPUT_TITLE)
            print_verbose(f'\tTemplate file {row.template_file}.', verbose=self.output_verbose, end='...',
                          **OUTPUT_TEXT)
            try:
                if row.template_file not in invalid:
                    fp_template = resolve_file_path(row.template_file)
                    self.batch_df.at[row.Index, 'template_file'] = fp_template
                    print_verbose(f"{self.batch_df.at[row.Index, 'template_file']}", verbose=self.output_verbose,
                                  **OUTPUT_TEXT)
            except ValueError as v:
                print_verbose(f'{v}: Row {row.Index} - Template file {row.template_file} does not exist at the given '
                              f'location.', True, **EXCEPTION_TEXT)
            except Exception as e:
                print_verbose(f'{e}: Row {row.Index}\n{row}', True, **EXCEPTION_TEXT)

            # Check 4.
            print_verbose(f'\tCheck output filename {row.output_file}.', verbose=self.output_verbose, end='...',
                          **OUTPUT_TEXT)
            if str(row.output_file) in invalid:
                raise ValueError(f'The name of the output file has not been provided.')
            try:
                fp_name = str(Path(row.output_file).name)
                fp_output = Path(resolve_file_path(Path(row.output_file).parent)).joinpath(fp_name)
                self.batch_df.at[row.Index, 'output_file'] = fp_output
                print_verbose(f"{self.batch_df.at[row.Index, 'output_file']}", verbose=self.output_verbose,
                              **OUTPUT_TEXT)
            except FileNotFoundError as f:
                print_verbose(f'{f}.  Row {row.Index} - Output file {row.output_file} directory does not exist at the '
                              f'given loaction.', True, **EXCEPTION_TEXT)
            except Exception as e:
                print_verbose(f'{e}: Row {row.Index}\n{row}', True, **EXCEPTION_TEXT)

            # Check 5.
            print_verbose(f'\tCheck row filters:', end='...', verbose=self.output_verbose, **OUTPUT_TEXT)
            if str(row.filter_rows).lower() not in invalid and row.filter_rows is not None:
                self.batch_df.at[row.Index, 'filter_rows'] = self.prepare_row_filters(row.filter_rows)
                print_verbose(f'Ok', verbose=self.output_verbose, **OUTPUT_TEXT)
            else:
                print_verbose(f'Ok. No filters exist.', verbose=self.output_verbose, **OUTPUT_TEXT)

            # Check 6.
            print_verbose(f'\tCheck drop empty columns', verbose=self.output_verbose, end='...', **OUTPUT_TEXT)
            if row.drop_empty_columns is None:
                self.batch_df.at[row.Index, 'drop_empty_columns'] = False
                print_verbose(f'Ok', verbose=self.output_verbose, **OUTPUT_TEXT)
            else:
                print_verbose(f'Ok', verbose=self.output_verbose, **OUTPUT_TEXT)

            # Check 7
            print_verbose(f'\tCheck header row details', verbose=self.output_verbose, end='...', **OUTPUT_TEXT)
            if row.header_row is None:
                self.batch_df.at[row.Index, 'header_row'] = 0
            print_verbose(f'Ok', verbose=self.output_verbose, **OUTPUT_TEXT)

    def check_structure_worksheet_data(self):
        """
        Check 1: Confirm expected batch headers exist in the batch worksheet.
        Check 2: Confirm the expected section types exist.
        Check 3: Confirm the section_contains values exist in the structure document.
        Check 4: Check the photo file paths and resolve
        Check 5: Check if section_break is None, set it to False.
        Check 6: Check if page_break is None, set it to False.
        :return:
        """
        t_structure_headers = list(self.t_structure_df)

        # Check 1
        self.data_check(f'  Check structure work sheet headers', f'Ok',
                        [(EXPECTED_STRUCTURE_HEADERS, t_structure_headers)], compare='subset')
        t_structure_section_types = list(self.t_structure_df.loc[:, 'section_type'])
        t_data_section_types = list(self.t_data_df)

        # Check 2
        t_structure_section_contains: List = []
        for each in list(self.t_structure_df.loc[:, 'section_contains']):
            for item in split_str(each):
                t_structure_section_contains.append(item)
        self.data_check(f'  Check structure worksheet section_types are correct.', f'Ok',
                        [(EXPECTED_SECTION_TYPES, t_structure_section_types)], compare='part')

        # Check 3
        self.data_check(f'  Check structure worksheet section_contain details are correct.', f'Ok',
                        [(t_data_section_types, t_structure_section_contains)], compare='part')

        for row in self.t_structure_df.itertuples():
            # Check 4
            if 'photo' == str(row.section_type).lower():
                root = self._input_fp.parent
                try:
                    print_verbose(f'  Row {row.Index}: Check file {row.path}', verbose=self.output_verbose,
                                  **OUTPUT_TEXT)
                    t_photo_path = resolve_file_path(root.joinpath(row.path))
                    if t_photo_path.is_dir():
                        self.batch_df.at[row.Index, 'path'] = t_photo_path
                        # Add the photo path to the dictionary with the section_contains as the key
                        self.t_structure_photo_path[row.section_contains] = t_photo_path
                    print_verbose(f'\t{self.t_structure_photo_path[row.section_contains]}', verbose=self.output_verbose,
                                  **OUTPUT_TEXT)

                except ValueError as v:
                    print_verbose(f'{v}: Row {row.Index} - Photo directory {row.path} cannot be found at the given '
                                  f'locaton.', True, **EXCEPTION_TEXT)
                except Exception as e:
                    print_verbose(f'{e}: Row {row.Index}\n{row}', True, **EXCEPTION_TEXT)

            # Check 5.
            print_verbose(f'\tCheck section break details', verbose=self.output_verbose, end='...', **OUTPUT_TEXT)
            if row.section_break is None:
                self.batch_df.at[row.Index, 'section_break'] = False
            print_verbose(f'Ok', verbose=self.output_verbose, **OUTPUT_TEXT)

            # Check 6.
            print_verbose(f'\tCheck page break details', verbose=self.output_verbose, end='...', **OUTPUT_TEXT)
            if row.page_break is None:
                self.batch_df.at[row.Index, 'page_break'] = False
            print_verbose(f'Ok', verbose=self.output_verbose, **OUTPUT_TEXT)

    def check_data_worksheet_data(self):
        """
        Check 1. Check that the photos exist in the directory. The check assumes that the first file name with the same
        name is the correct file if no filename has been provided in the worksheet. The check assumes that all photo
        names are unique regardless of photo directory.
        :return:
        """
        # Check 1. check the photo paths.
        t_photos_found: Dict[str, Path] = {}
        # Assuming that that there may be more than one directory containing photos for the worksheet loop through the
        # each folder.
        columns = self.t_structure_photo_path.keys()
        print_verbose(f'  Photos are located in the following data worksheet columns:', verbose=self.output_verbose,
                      **OUTPUT_TEXT)
        t_columns = enumerate(columns, 1)
        for item, col in t_columns:
            print_verbose(f'\t{item}\t{col}', verbose=self.output_verbose, **OUTPUT_TEXT)

        # Grab the photos stored in the folders and store their paths in dictionary. Store them in the dictionary
        # with their name minus the file extension as the key.
        for col in columns:
            for file_ext in photo_formats:
                for file in self.t_structure_photo_path[col].glob('*' + file_ext):
                    t_photos_found[file.name] = file
            # For each of the columns containing photos loop through the self.t_data_df and replace the file name with
            # the path.
            t_df_rows = self.t_data_df.itertuples()
            for row in t_df_rows:
                # The row below must remain in this position
                col = str(col).lower()
                t_row = row._asdict()
                t_row_photo = []
                print_verbose(f'  Row {row.Index}:', verbose=self.output_verbose, **OUTPUT_TITLE)
                if str(t_row[col]).lower() not in ['no photo', 'none', 'nan', '-']:
                    for t in split_str(t_row[col]):
                        try:
                            t_row_photo.append(self.check_photo_paths(t, t_photos_found))
                        except Exception as e:
                            print_verbose(f'{e}: Row {row.Index} - {t_row[col]}', True, **EXCEPTION_TEXT)
                        self.t_data_df.at[t_row['Index'], col] = t_row_photo
                if isinstance(self.t_data_df.at[t_row['Index'], col], Iterable):
                    for fp in self.t_data_df.at[t_row['Index'], col]:
                        if isinstance(fp, Path):
                            print_verbose(f'\t{str(fp)}', verbose=self.output_verbose, **OUTPUT_TEXT)

    @staticmethod
    def check_photo_paths(expected_photo: (Path, str), actual_photos: dict) -> Path:
        """
        Check 1. Check that the photo could be an image file.
        Check 2. If the expected_photo does not have a file extension then loop through photo_formats and see if the
        file does exist.
        Check 3. If the expected_photo does have a file extension then check that the file does exist.
        The method will raise exceptions if the file is not found
        :param expected_photo:
        :param actual_photos:
        :return:
        """
        # Check 1
        photo = expected_photo.strip()
        if Path(photo).suffix is not '' and Path(photo).suffix not in photo_formats:
            raise ValueError(f'The data worksheet photo {photo} is not been specified as a photo. Ensure that the'
                             f' file format is one of the following formats {photo_formats}.')
        # Check 2
        elif Path(photo).suffix is '':
            for ext in photo_formats:
                try:
                    t_photo_name = str(photo + ext)
                    if actual_photos[t_photo_name]:
                        return Path(actual_photos[t_photo_name]).resolve(strict=True)
                except KeyError as k:
                    print_verbose(f'{k}', True, **EXCEPTION_TEXT)
            raise ValueError(f'The photo {photo} does not exist in the specified directory.')
        # Check 3
        elif Path(photo).suffix in photo_formats:
            try:
                if actual_photos[photo]:
                    return Path(actual_photos[photo]).resolve(strict=True)
            except KeyError as k:
                print_verbose(f'{k}. The photo {photo} does not exist in the directory.', True, **EXCEPTION_TEXT)

        raise ValueError(f'The photo {photo} does not appear to exist in the directory.')

    @staticmethod
    def excel_to_dataframe(io, worksheet: str, header_row: int = 0, clean_header: bool = False,
                           drop_empty_rows: bool = False) -> data_frame:
        """
        Open and perform basic cell_data cleaning on a single excel work worksheet.
        :param io: The Excel file to be read.
        :param worksheet: The Excel spreadsheet worksheet's name.
        :param header_row: index of the header row in the spreadsheet. Defaults to 0, i.e. assumes the headers are at
        the top of the page.
        :param clean_header: If True clean the column headers
        :param drop_empty_rows: If True remove empty rows.
        :return:
        """
        df = pd.read_excel(io, worksheet, header_row)
        if clean_header is not False:
            try:
                df = df.clean_names()
            except KeyError as k:
                print_verbose(f'{k}', True, **EXCEPTION_TEXT)
        if drop_empty_rows is True:
            try:
                df = df.dropna(thresh=2)
            except KeyError as k:
                print_verbose(f'{k}', True, **EXCEPTION_TEXT)
        return df

    @staticmethod
    def prepare_row_filters(filters: str) -> List[Tuple[Any, list]]:
        """
        Split the passed string into a list of Tuples taking the form (column_name, filter_keyword)
        :param filters:
        :return:
        """
        filtered_list = []
        i = []
        if '\n' in filters:
            i = str(filters).splitlines()
        else:
            i.append(filters)
        for each in i:
            col, b = each.split(':')
            col_kw = strip_whitespace(b.split(','))
            filtered_list.append((col, col_kw))
        return filtered_list

    @staticmethod
    def compare_lists(expected_list, actual_list):
        expected_list = sorted(expected_list)
        actual_list = sorted(actual_list)
        if set(expected_list).issubset(set(actual_list)) is False:
            raise ValueError(f'The provided headers:\n\t{actual_list}\ndo not match the required headers'
                             f'\n\t{expected_list}.')
        else:
            return True

    @staticmethod
    def in_lists(expected_list, actual_list):
        for each in actual_list:
            if each.lower() not in expected_list:
                raise ValueError(f'The provided headers:\n\t{actual_list}\ndo not match the required headers'
                                 f'\n\t{expected_list}.')
        else:
            return True

    def data_check(self, check_text: str, success_text: str, comparision_list: List[tuple], compare: str):
        """
        :param check_text:
        :param success_text:
        :param comparision_list:
        :param compare:
        :return:
        """
        try:
            print_verbose(f'{check_text}', verbose=self.output_verbose, end='...', **OUTPUT_TEXT)
            if compare is 'subset':
                for expected, actual in comparision_list:
                    self.compare_lists(expected, actual)
            if compare is 'part':
                for expected, actual in comparision_list:
                    self.in_lists(expected, actual)
            print_verbose(f'{success_text}', verbose=self.output_verbose, **OUTPUT_TEXT)
        except ValueError as v:
            print_verbose(f'\nValueError:\n{v}', True, **EXCEPTION_TEXT)
            exit_app()
        except Exception as e:
            print_verbose(f'General exception {e}', True, **EXCEPTION_TEXT)
            exit_app()

    def check_dataframe(self, title: str, check_dataframe: pd.DataFrame, worksht_title: str, check_method,
                        complete_check: str, check_worksheet: str = '', exception_text: str = ''):
        try:
            print_verbose(f'{title}: {check_worksheet}', verbose=self.output_verbose, **DATAFRAME_TITLE)
            print_verbose(f'{check_dataframe}', verbose=self.output_verbose, **DATAFRAME_TEXT)
            print_verbose(f'{worksht_title}: {check_worksheet}', verbose=self.output_verbose, **OUTPUT_TITLE)
            check_method()
            print_verbose(f'{complete_check}: {check_worksheet}', verbose=self.output_verbose, **OUTPUT_TITLE)
        except KeyError as k:
            print_verbose(f'KeyError {k}: ', True, **EXCEPTION_TEXT)
            exit_app()
        except ValueError as v:
            print_verbose(f'\nValueError {v}: \n', True, **EXCEPTION_TEXT)
            exit_app()
        except Exception as e:
            print_verbose(f'{e}: ', True, **EXCEPTION_TEXT)
            exit_app()
