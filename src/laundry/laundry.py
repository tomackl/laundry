from docx import Document
from docx.shared import Inches
import pandas as pd
import janitor
from pathlib import Path, PurePath
from typing import List, Iterable, Dict, Tuple, Any, NewType, Iterator
import click

laundry_version = '2019.0.7'

data_frame = NewType('data_frame', pd.DataFrame)
invalid = ['nan', 'None', 'NA', 'N/A', 'False', 'Nil']


def clean_xlsx_table(file_path: str, sheet: str, head: int = 0,
                     rm_column: List[str] = None, clean_hdr: bool = False,
                     drop_empty: bool = False) -> data_frame:
    """
    Open and perform basic data cleaning on a single excel work sheet.
    :param file_path: path to the excel file
    :param sheet: excel spreadsheet name
    :param head: index of the header row in the spreadsheet. Defaults to 0.
    :param rm_column: remove the column headers contained in the passed list.
    :param clean_hdr: If True clean the column headers
    :param drop_empty: If True remove empty rows.
    :return:
    """
    # todo: add exception here to catch a XLRDError in the event of a misnamed work sheet.
    df = pd.read_excel(pd.ExcelFile(file_path), sheet, head)
    if rm_column is not None:
        df = df.remove_columns(rm_column)
    if clean_hdr is not False:
        df = df.clean_names()
    if drop_empty is True:
        df = df.dropna(thresh=2)
    return df


def extract_data(record: Dict, header: List[str], format_title: bool = True
                 ) -> List[Tuple]:
    """
    Take a dictionary and split in to a list of tuples containing the 'keys'
    data defined in 'header' as the first tuple and the associated values as the second.
    The function will return a list containing two equal length tuples.
    :param record: the dictionary containing the data.
    :param header: the keys that defined the key-values to be extracted.
    :param format_title: make the header string title case.
    :return: a list of tuples
    """
    hdr, data = [], []
    for each in header:
        h = remove_underscore(each)
        if format_title is True:
            hdr.append(str(h).title())
        else:
            hdr.append(str(h))
        t = record.pop(each)
        data.append(t)
    return [tuple(hdr), tuple(data)]


def remove_underscore(text: str) -> str:
    return text.replace('_', ' ')


def section_contains(sect_contains: Any) -> List[str]:
    """
    A function to split a string into a list. It will return a list.
    :param sect_contains:
    :return:
    """
    if '\n' in sect_contains:
        return list(str(sect_contains).splitlines())
    elif ',' in sect_contains:
        return list(str(sect_contains).split(','))
    else:
        i = list()
        i.append(sect_contains)
        return i


def insert_paragraph(document: object, text: str, title: str = None,
                     section_style: str = None, title_style: str = None,
                     ):
    """
    This function defines the characteristics of a paragraph to be added to the document.
    :param document: the template_doc the paragraph will be added to.
    :param text: paragraph text
    :param title: title text
    :param section_style: paragraph style
    :param title_style: title style. Use this _or_ title_level.
    :return:
    """
    if text == '':
        document.add_paragraph(text)
        return
    text = text.splitlines()
    if (title is not None) and (str(title_style) != 'nan'):
        document.add_paragraph(remove_underscore(title), style=title_style)
    for each in text:
        document.add_paragraph(each, style=section_style)


def insert_table(document: object, cols: int, rows: int,
                 data: List[Iterable[str]], section_style: str = None,
                 autofit: bool = True):
    """
    The function takes data and uses it to create a table for the template_doc.
    The first row of data is assumed to be the table header.
    20190814 'add_table_hdr added to allow the table header to be dropped from the table if not required.
    :param document: the docx file the table will be added to.
    :param rows: the number of required table rows.
    :param cols: the number of required table columns.
    :param data: The list data to be inserted into the table. The idx[0] is assumed to be the header.
    :param section_style: The style to be used for the table.
    :param autofit: autofit the table to the page width.
    :return:
    """
    # todo: add error checking for 'cols' and 'row'
    table: object = document.add_table(rows=rows, cols=cols, style=section_style)
    if autofit is True:
        table.autofit = True
    data = enumerate(data, 0)
    for i, cell_contents in data:
        insert_row(table.rows[i].cells, cell_contents)


def insert_row(row_cells, data: List[str]) -> Any:
    """
    Populate a table row. The cells are passed as a row and the contents added.
    :param row_cells:
    :param data:
    :param style: style is the text style to be applied to the individual rows.
    :return:
    """
    for i, text in enumerate(data):
        row_cells[i].text = str(text)
    return row_cells


def insert_photo(document: object, photo: str, width: int = 4):
    """
    Insert a photo located at path into document and set the photo to width.
    :param document: the docx file the photo will be added to.
    :param photo: file path to the
    :param width: width of the image in Inches
    :return:
    """
    photo_formats = ['', '.jpg', '.jpeg', '.png', '.tiff']
    for ext in photo_formats:
        photo_path = Path(photo + ext)
        if photo_path.exists():
            document.add_picture(str(photo_path), width=Inches(width))
            return True
    print(f'Photo {photo.lower()} does not exist. Check file extension.')
    document.add_paragraph(f'PHOTO "{str(photo).upper()}" NOT FOUND\n')


def format_docx(rowdict: dict, structdict: dict, outputfile: object, file_path: str):
    """
    The function is passed a dict (data_dict) containing the data to be formatted
    (structure) based on the template (outputfile).
    :param rowdict: dictionary containing the text. It represents a single row from the spreadsheet.
    :param structdict: defines the output file's format structure.
    :param outputfile: The file which data will be inserted into.
    :param file_path: directory containing the spreadsheet
    :return:
    """
    file_path = file_path
    # todo: add error checking here.
    for element in structdict:
        element_sect_con = str(element['sectioncontains']).lower()
        if str(element['sectiontype']).lower() in ('heading', 'para', 'paragraph'):
            insert_paragraph(outputfile, str(rowdict[str(element_sect_con).lower()]),
                             title=element_sect_con.title(),
                             section_style=element['sectionstyle'],
                             title_style=element['titlestyle']
                             )

        elif str(element['sectiontype']).lower() == 'table':
            table = section_contains(element_sect_con)
            data = extract_data(rowdict, table)
            insert_table(outputfile, len(table), len(data),
                         data, section_style=element['sectionstyle']
                         )

        elif str(element['sectiontype']).lower() == 'photo':
            sect_contains = rowdict[element_sect_con]
            q = confirm_path_directory([file_path, element['path']])
            if str(sect_contains).lower() not in ['no photo', 'none', 'nan', '-']:
                photo = section_contains(sect_contains)
                for each in photo:
                    loc = q.joinpath(each)
                    insert_photo(outputfile, str(loc), 4)

        else:
            print('Valid section header was not found.')

        if element['sectionbreak'] is True:
            insert_paragraph(outputfile, '')

        if element['pagebreak'] is True:
            outputfile.add_page_break()


def confirm_path_directory(filepath: List[str]) -> Path:
    """
    Convert the contents of the passed list into a Path. This function assumes
    that the sum of the passed list will be a single path to a directory.
    :param filepath: a list of path names as string
    :return: Path
    """
    filepath = filepath
    p = PurePath()
    for each in filepath:
        try:
            q = PurePath(each.replace('\\', '/').strip('/'))
        except Exception as e:
            print(f'{e}:\nPath element "{each}" is not a valid path name.')
            # False
        p = p / Path(q.as_posix())
    r = Path(p)
    if r.is_dir():
        return r
    return False


def confirm_path_file(filepath: List[str]) -> bool:
    """
    Convert the contents of the passed list into a Path and if it points to a file
    return True.
    :param filepath:
    :return:
    """
    filepath = filepath
    p = PurePath()
    for each in filepath:
        q = PurePath(each.replace('\\', '/').strip('/'))
        p = p / Path(q.as_posix())
    p = Path(p)
    return p.exists()


def worksheet_present(sheet_names: List[str], sheets: List[str]) -> bool:
    """
    Check whether the worksheets in sheets exist within the spreadsheet returning
    True if they are present.
    :param sheet_names:
    :param sheets:
    :return:
    """
    return set(sheet_names) >= set(sheets)


def remove_columns(load: data_frame, columns: List[str]) -> dict:
    """
    Pass a dataframe dictionary and list with the column names to be dropped.
    :param load:
    :param columns:
    :return:
    """
    return load.drop(columns=columns)


def single_load(structure_dict: Dict, data_dict: Dict, file_template: str, path_input_f: str,
                file_output: str):
    """
    This function controls the production of the output file and is called for both the single and multi modes.
    :param structure_dict: defines the structure of the output file
    :param data_dict: contains the data to be manipulated and exported in the output file.
    :param file_template: template document that will form the basis of the output file.
    :param path_input_f: path to the current working directory.
    :param file_output: output file name.
    """
    with click.progressbar(iterable=data_dict,
                           label='Conversion progress:',
                           fill_char='|',
                           empty_char='_'
                           ) as data_dictionary:
        for row in data_dictionary:
            format_docx(row, structure_dict, file_template, file_path=str(path_input_f))
    try:
        file_template.save(file_output)
    except FileNotFoundError as e:
        p = Path(file_output)
        print(f'{e}:\nCheck that output directory "{p.resolve().parent}" exists.')
        print(f'Check your spelling. ;)')
    except Exception as e:
        print(f'{e}')


# def confirm_path(file_path: str, rel_path=True) -> bool:
#     """
#     The function returns True if the path exists and false if doesn't. The function assumes that the paths are relavtive
#     and resolves from CWD.
#     :param file_path:
#     :return:
#     """
#     p = Path(file_path)
#     print(p)
#     # p = Path(str(Path.cwd()) + '/' + file_path)
#     if rel_path is True and p.exists() is True:
#     # if p.exists() is True:
#         return True
#     if rel_path is False and p.parent.exists() is True:
#         return True
#     else:
#         print(f"Path '{file_path}' is referenced in the worksheet but cannot be found.")
#         return False


@click.group()
@click.version_option(laundry_version)
def cli():
    """
    This is the command line interface(CLI) for the Laundry app.
    For details regarding the operation of the app type `laundry --help`.
    """
    pass


@cli.command()
@click.option('--data-worksheet', '-dw', 'data',
              default='Master List',
              help='Name of the worksheet containing the data to be converted into a '
                   'word document. '
                   'The default is "Master List".'
              )
@click.option('--template', '-t', 'template',
              help='Name of the template file to be used used as the basis of the '
                   'converted file.',
              type=click.Path(exists=True)
              )
@click.option('--structure-worksheet', '-sw', '-s', 'structure',
              default='_structure',
              help='Name of the worksheet containing the data to format the structure '
                   'of the outfile document. The default is "_structure".'
              )
@click.option('--data-header', '-dh', 'data_head',
              default=0,
              type=int,
              help="The row number of the data worksheet's row containing the column "
                   "headers. The default is 0."
              )
@click.argument('input_file',
                type=click.Path(exists=True)
                )
@click.argument('output_file')
def single(input_file, output_file, data, structure, template, data_head):
    """
    Run laundry on a single worksheet.

    The relative path for each file should be provided with each of the options if non-default file names are provided.

    NOTE: If output files are intended to be saved in a separate directory, that directory *must* exist otherwise the
    output file will not save.

    IMPORTANT: Laundry will overwrite, without prompting, any files with the same name in the directory where output
    files are saved.
    """
    file_input = Path(input_file)
    file_output = output_file
    wkst_data = data
    wkst_struct = structure
    template = template
    wash_single(file_input, file_output, wkst_data, wkst_struct, template, data_head)


@cli.command()
@click.option('--batch-worksheet', '-b', 'batch',
              default='_batch',
              help='Name of the worksheet containing the format data. This '
                   'worksheet defines the structure and data worksheets and '
                   'other higher level formatting details. The default batch '
                   'worksheet name is "_batch".')
@click.argument('input_file',
                type=click.Path(exists=True)
                )
def multi(input_file, batch):
    """
    Run Laundry on multiple worksheets.
    """
    file_input = Path(input_file)
    wksht_batch = batch
    wash_multi(file_input, wksht_batch)


def wash_single(file_input, file_output, wksht_data, wksht_struct, template, data_head):
    """
    This function acts as a common calling point for the module to allow the module to be run from the command line
    interface (cli) or from another script.
    :param file_input: the .xls file containing the data to be converted.
    :param file_output: name of the output file.
    :param wksht_data: name of the .xls worksheet containing the data to be processed
    :param wksht_struct: name of the .xls worksheet detailing how the data shall be processed
    :param template: the .docx file to be used as the template
    :param data_head: the number of the data worksheet's row containing the column headers.
    """
    # todo: add exception to ensure that the `template` file actually exists.
    #  `docx.opc.exceptions.PackageNotFoundError' is raised if the file does not exist.
    file_template = Document(template)
    # todo: add exceptions to catch files that are missing file extensions.
    path_input_f = file_input.parents[0]
    check_load = pd.ExcelFile(file_input).sheet_names
    if worksheet_present(check_load, [wksht_struct, wksht_data]):
        structure_file = clean_xlsx_table(file_input, sheet=wksht_struct, head=0,
                                          clean_hdr=True, drop_empty=False
                                          )
        data_file = clean_xlsx_table(file_input, sheet=wksht_data, head=data_head,
                                     clean_hdr=True, drop_empty=True
                                     )
        single_load(structure_file.to_dict('records'), data_file.to_dict('records'),
                    file_template, path_input_f, file_output)
    else:
        no_valid_data('Valid worksheets were not found.', [wksht_struct, wksht_data],
                      'Required worksheets', check_load, 'Found worksheets')


def wash_multi(file_input, wksht_batch):
    """
    This function acts as a common calling point for the module to allow the module to be run from the command line
    interface (cli) or from another script.
    :param file_input: the .xls file containing the data to be converted.
    :param wksht_batch: name of the .xls worksheet detailing how the data shall be processed
    """
    # todo: add exception to ensure that the `template` file actually exists.
    #  `docx.opc.exceptions.PackageNotFoundError' is raised if the file does not exist.
    # todo: add exceptions to catch files that are missing file extensions.
    path_input_f = file_input.parents[0]

    check_load = pd.ExcelFile(file_input).sheet_names
    if worksheet_present(check_load, [wksht_batch]):

        format_file = clean_xlsx_table(file_input, sheet=wksht_batch, head=0,
                                       clean_hdr=True, drop_empty=False,
                                       )
        sort_colours(format_file.to_dict('records'), check_load, file_input, path_input_f)

    else:
        no_valid_data('Valid worksheets were not found.', wksht_batch,
                      'Required worksheets', check_load, 'Found worksheets')


def no_valid_data(descrip, req_data, req_title, actual_data, actual_title):
    """
    Output to the user the data that was required and what was provided.
    :param descrip: Description of the no_valid_data
    :param req_data: The data that was expected
    :param req_title: What is te required data called?
    :param actual_data: The data that was found
    :param actual_title: What is the actual data called>?
    :return:
    """
    print(f'{descrip}')
    print(f'\t{req_title} -> {req_data}')
    print(f'\t{actual_title} -> {actual_data}')


def sort_colours(load: Dict, check_load, file_input, path_input_f):
    """
    This function will control take the _format worksheet and call the appropriate
    functions to convert the files.
    :param load: the 'batch' dictionary,
    :param check_load: a list containing the spreadsheets
    :param file_input:
    :param path_input_f:
    :return:
    """

    for row in load:
        if not worksheet_present(check_load, [row['structure_worksheet'], row['data_worksheet']]):
            no_valid_data('The following worksheet was missing.', row['structure_worksheet'],
                          'Structure worksheet.', row['data_worksheet'], 'Data worksheet')
            break
        elif not confirm_path_file([row['template_file']]):
            no_valid_data('Template file not found', row['template_file'], 'File path passed',
                          Path(str(row['template_file'])).resolve(), 'Resolved template file path.')
            break
        sf = clean_xlsx_table(file_input,
                              row['structure_worksheet'],
                              head=0,
                              clean_hdr=True,
                              drop_empty=False
                              )
        df = clean_xlsx_table(file_input,
                              row['data_worksheet'],
                              head=row['header_row'],
                              # rm_column=remove_columns,
                              clean_hdr=True,
                              drop_empty=True
                              )
        if pd.notna(row['filter_rows']):
            df = filter_rows(df, filter_setup(row['filter_rows']))
        single_load(sf.to_dict('records'), df.to_dict('records'),
                    Document(row['template_file']), path_input_f, str(row['output_file']))


def filter_rows(df: data_frame, filter_list: List[List[str]]) -> data_frame:
    """
    Takes a dataframe and returns another dataframe that only contains rows that meet filter_list.
    Do not the output the rows that meet certain conditions.
    Only can be selected in the _batch worksheet
    Takes the form <column_header>:<value1>, <value2>, ..., <valueN>
    - Multiple columns can be added
    - Multiple values for a given column can be provided
    - Conditional filtering is not provided -> if a column header and value is provided then the row will be filtered.
    :return:
    :param df: the dataframe to be filtered.
    :param filter_list: made up of <column>, (<row_value1>, ..., <row_valueN>).
    :return:
    """
    for col in filter_list:
        df = df.loc[df[col[0]].isin(col[1])]
    return df


def filter_setup(filters: str) -> List[List[str]]:
    """
    Split the content of the passed tuple into the correct parts
    :return:
    """
    filtered_list = []
    i = []
    try:
        if '\n' in filters:
            i = str(filters).splitlines()
        else:
            i.append(filters)
    except TypeError as e:
        print(f'The filter entered is: {filters}.\n')
        print(f'Row filtering must take the following form:')
        print(f'\t[ColumnHeader]:[CellContent] where:')
        print(f'\t[ColumnHeader] is the column heading text in lower case with underscores ("_" replacing spaces (" ").')
        print(f'\t[CellContent] is the text inside the cell that you want to find.')

    for each in i:
        column_hdr, filter_terms = each.split(':')
        column_hdr = column_hdr.lower().replace(' ', '_')
        filter_list = strip_list_whitespace(filter_terms.split(','))
        filtered_list.append((column_hdr, filter_list))
    return filtered_list


def strip_list_whitespace(wht_spc: List) -> List:
    i = []
    for each in wht_spc:
        i.append(each.strip())
    return i
